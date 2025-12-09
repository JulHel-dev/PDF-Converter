"""
PDF Converter

Handles PDF conversions to and from various formats.
Supported:
- PDF → TXT, MD, DOCX, HTML, PNG, JPEG
- Images → PDF
- Documents → PDF
"""
import os
from typing import Optional, Dict, List
from src.converters.base_converter import BaseConverter
from src.detection.ocr_detector import OCRDetector


class PDFConverter(BaseConverter):
    """
    Handles PDF conversions.
    
    Supported:
    - PDF → TXT, MD, DOCX, HTML, PNG, JPEG
    - Images → PDF
    """
    
    def __init__(self):
        super().__init__()
        self.supported_input_formats = ['pdf']
        self.supported_output_formats = ['txt', 'md', 'docx', 'html', 'png', 'jpeg', 'jpg']
        self.ocr_detector = OCRDetector()
    
    def convert(self, input_path: str, output_path: str, output_format: str) -> bool:
        """Convert PDF to target format."""
        timer_id = self.monitor.start_timer('pdf_conversion')
        
        try:
            # Step 1: Validate input
            self.monitor.log_event('conversion_start', {
                'input': input_path, 
                'output_format': output_format,
                'output_path': output_path
            }, severity='INFO', step='Step 1/5: Validating input')
            
            if not self.validate_input(input_path):
                return False
            
            if not self.validate_output(output_path):
                return False
            
            # Step 2: Check text layer
            self.monitor.log_event('text_layer_check', {
                'input': input_path
            }, severity='INFO', step='Step 2/5: Checking text layer')
            
            has_text = self.ocr_detector.check_pdf_text_layer(input_path)
            
            if not has_text:
                self.monitor.log_event('no_text_layer', {
                    'input': input_path,
                    'recommendation': 'Run OCR externally (e.g., Foxit Compressor, Adobe Acrobat)'
                }, severity='WARNING', step='Step 2/5: No text layer detected')
            
            # Step 3: Extract content
            self.monitor.log_event('content_extraction', {
                'input': input_path, 
                'has_text_layer': has_text
            }, severity='INFO', step='Step 3/5: Extracting content')
            
            content = self._extract_content(input_path, has_text)
            
            if not content:
                self.monitor.log_event('extraction_failed', {
                    'input': input_path
                }, severity='ERROR', step='Step 3/5: Content extraction failed')
                return False
            
            # Step 4: Convert to target format
            self.monitor.log_event('format_conversion', {
                'output_format': output_format
            }, severity='INFO', step='Step 4/5: Converting to target format')
            
            result = self._write_output(content, output_path, output_format)
            
            # Step 5: Complete
            self.monitor.log_event('conversion_complete', {
                'input': input_path, 
                'output': output_path, 
                'success': result
            }, severity='INFO', step='Step 5/5: Conversion complete')
            
            return result
            
        except Exception as e:
            self.monitor.log_event('conversion_failed', {
                'input': input_path,
                'output_format': output_format,
                'error': str(e),
                'error_type': type(e).__name__
            }, severity='ERROR', step='Conversion failed')
            return False
            
        finally:
            self.monitor.end_timer(timer_id)
    
    def _extract_content(self, pdf_path: str, has_text: bool) -> Optional[Dict]:
        """Extract content from PDF with fallback for corrupt files."""
        try:
            import fitz  # PyMuPDF
            
            content = {
                'text': '',
                'pages': [],
                'metadata': {},
                'page_count': 0
            }
            
            with fitz.open(pdf_path) as doc:
                content['page_count'] = len(doc)
                content['metadata'] = doc.metadata
                
                # Extract text from each page
                for page_num, page in enumerate(doc):
                    try:
                        page_text = page.get_text('text')
                        content['pages'].append({
                            'number': page_num + 1,
                            'text': page_text,
                            'width': page.rect.width,
                            'height': page.rect.height
                        })
                        content['text'] += f"\n\n--- Page {page_num + 1} ---\n{page_text}"
                    except Exception as e:
                        # Partial extraction fallback
                        self.monitor.log_event('page_extraction_failed', {
                            'page': page_num + 1,
                            'error': str(e)
                        }, severity='WARNING')
                        content['pages'].append({
                            'number': page_num + 1,
                            'text': f'[Page {page_num + 1} extraction failed]',
                            'error': str(e)
                        })
            
            return content
            
        except ImportError:
            self.monitor.log_event('dependency_missing', {
                'library': 'PyMuPDF',
                'install_command': 'pip install PyMuPDF'
            }, severity='ERROR')
            return None
        except Exception as e:
            self.monitor.log_event('content_extraction_failed', {
                'file': pdf_path,
                'error': str(e)
            }, severity='ERROR')
            return None
    
    def _write_output(self, content: Dict, output_path: str, format: str) -> bool:
        """Write extracted content to target format."""
        format = format.lower()
        
        try:
            if format == 'txt':
                return self._write_text(content, output_path)
            elif format == 'md':
                return self._write_markdown(content, output_path)
            elif format == 'html':
                return self._write_html(content, output_path)
            elif format == 'docx':
                return self._write_docx(content, output_path)
            elif format in ['png', 'jpeg', 'jpg']:
                return self._write_images(content, output_path, format)
            else:
                self.monitor.log_event('unsupported_output_format', {
                    'format': format
                }, severity='ERROR')
                return False
                
        except Exception as e:
            self.monitor.log_event('output_write_failed', {
                'format': format,
                'error': str(e)
            }, severity='ERROR')
            return False
    
    def _write_text(self, content: Dict, output_path: str) -> bool:
        """Write content as plain text."""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content['text'])
        return True
    
    def _write_markdown(self, content: Dict, output_path: str) -> bool:
        """Write content as Markdown."""
        md_content = f"# Document\n\n"
        
        # Add metadata if available
        metadata = content.get('metadata', {})
        if metadata:
            md_content += "## Metadata\n\n"
            for key, value in metadata.items():
                if value:
                    md_content += f"- **{key}**: {value}\n"
            md_content += "\n"
        
        md_content += "## Content\n\n"
        md_content += content['text']
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        return True
    
    def _write_html(self, content: Dict, output_path: str) -> bool:
        """Write content as HTML."""
        html = "<!DOCTYPE html>\n<html>\n<head>\n"
        html += "<meta charset=\"UTF-8\">\n"
        html += "<title>PDF Conversion</title>\n"
        html += "<style>body { font-family: Arial, sans-serif; margin: 40px; }</style>\n"
        html += "</head>\n<body>\n"
        html += "<h1>Document</h1>\n"
        
        # Add metadata
        metadata = content.get('metadata', {})
        if metadata:
            html += "<h2>Metadata</h2>\n<ul>\n"
            for key, value in metadata.items():
                if value:
                    html += f"<li><strong>{key}:</strong> {value}</li>\n"
            html += "</ul>\n"
        
        html += "<h2>Content</h2>\n"
        html += f"<pre>{content['text']}</pre>\n"
        html += "</body>\n</html>"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
        return True
    
    def _write_docx(self, content: Dict, output_path: str) -> bool:
        """Write content as DOCX."""
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading('PDF Conversion', 0)
            
            # Add metadata
            metadata = content.get('metadata', {})
            if metadata:
                doc.add_heading('Metadata', level=1)
                for key, value in metadata.items():
                    if value:
                        doc.add_paragraph(f"{key}: {value}")
            
            # Add content
            doc.add_heading('Content', level=1)
            for paragraph in content['text'].split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            doc.save(output_path)
            return True
            
        except ImportError:
            self.monitor.log_event('dependency_missing', {
                'library': 'python-docx',
                'install_command': 'pip install python-docx'
            }, severity='ERROR')
            return False
        except Exception as e:
            self.monitor.log_event('docx_write_failed', {
                'error': str(e)
            }, severity='ERROR')
            return False
    
    def _write_images(self, content: Dict, output_path: str, format: str) -> bool:
        """Write PDF pages as images."""
        try:
            import fitz
            from src.config.settings import DEFAULT_DPI
            
            # For single page, use the output path directly
            # For multiple pages, create numbered files
            pdf_path = content.get('pdf_path')
            if not pdf_path:
                # Need to re-open PDF - extract path from somewhere
                self.monitor.log_event('image_conversion_warning', {
                    'reason': 'PDF path not in content, cannot render images'
                }, severity='WARNING')
                return False
            
            with fitz.open(pdf_path) as doc:
                page_count = len(doc)
                
                if page_count == 1:
                    # Single page - use output path directly
                    page = doc[0]
                    pix = page.get_pixmap(dpi=DEFAULT_DPI)
                    pix.save(output_path)
                else:
                    # Multiple pages - create numbered files
                    base_name = os.path.splitext(output_path)[0]
                    for page_num, page in enumerate(doc):
                        page_output = f"{base_name}_page_{page_num + 1}.{format}"
                        pix = page.get_pixmap(dpi=DEFAULT_DPI)
                        pix.save(page_output)
            
            return True
            
        except ImportError:
            self.monitor.log_event('dependency_missing', {
                'library': 'PyMuPDF',
                'install_command': 'pip install PyMuPDF'
            }, severity='ERROR')
            return False
        except Exception as e:
            self.monitor.log_event('image_write_failed', {
                'error': str(e)
            }, severity='ERROR')
            return False
    
    def pdf_to_text(self, pdf_path: str) -> str:
        """
        Extract all text from PDF.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            import fitz
            
            with fitz.open(pdf_path) as doc:
                return '\n\n'.join(
                    f"--- Page {i+1} ---\n{page.get_text('text')}" 
                    for i, page in enumerate(doc)
                )
        except Exception as e:
            self.monitor.log_event('text_extraction_failed', {
                'file': pdf_path,
                'error': str(e)
            }, severity='ERROR')
            return ""
    
    def pdf_to_images(self, pdf_path: str, output_dir: str, dpi: int = 150) -> List[str]:
        """
        Render each PDF page as an image.
        
        Args:
            pdf_path: Path to PDF file
            output_dir: Directory for output images
            dpi: Resolution for rendering
            
        Returns:
            List of created image file paths
        """
        try:
            import fitz
            
            os.makedirs(output_dir, exist_ok=True)
            image_paths = []
            
            with fitz.open(pdf_path) as doc:
                for page_num, page in enumerate(doc):
                    output_path = os.path.join(output_dir, f"page_{page_num + 1}.png")
                    pix = page.get_pixmap(dpi=dpi)
                    pix.save(output_path)
                    image_paths.append(output_path)
            
            return image_paths
            
        except Exception as e:
            self.monitor.log_event('image_rendering_failed', {
                'file': pdf_path,
                'error': str(e)
            }, severity='ERROR')
            return []
