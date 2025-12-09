"""
Markdown Converter

Handles Markdown conversions.
Supported: MD ↔ HTML, TXT, PDF, DOCX
"""
import os
from typing import Optional, Dict
from src.converters.base_converter import BaseConverter


class MarkdownConverter(BaseConverter):
    """
    Handles Markdown conversions.
    
    Supported: MD ↔ HTML, TXT, PDF, DOCX
    """
    
    def __init__(self):
        super().__init__()
        self.supported_input_formats = ['md', 'markdown']
        self.supported_output_formats = ['html', 'txt', 'pdf', 'docx']
    
    def convert(self, input_path: str, output_path: str, output_format: str) -> bool:
        """Convert Markdown to target format."""
        timer_id = self.monitor.start_timer('markdown_conversion')
        
        try:
            # Step 1: Validate
            self.monitor.log_event('conversion_start', {
                'input': input_path,
                'output_format': output_format,
                'output_path': output_path
            }, severity='INFO', step='Step 1/4: Validating input')
            
            if not self.validate_input(input_path):
                return False
            
            if not self.validate_output(output_path):
                return False
            
            # Step 2: Load markdown
            self.monitor.log_event('markdown_load', {
                'input': input_path
            }, severity='INFO', step='Step 2/4: Loading markdown')
            
            with open(input_path, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            # Step 3: Convert to target format
            self.monitor.log_event('format_conversion', {
                'output_format': output_format
            }, severity='INFO', step='Step 3/4: Converting to target format')
            
            result = self._write_output(markdown_text, output_path, output_format)
            
            # Step 4: Complete
            self.monitor.log_event('conversion_complete', {
                'input': input_path,
                'output': output_path,
                'success': result
            }, severity='INFO', step='Step 4/4: Conversion complete')
            
            return result
            
        except Exception as e:
            self.monitor.log_event('conversion_failed', {
                'input': input_path,
                'output_format': output_format,
                'error': str(e),
                'error_type': type(e).__name__
            }, severity='ERROR')
            return False
        finally:
            self.monitor.end_timer(timer_id)
    
    def _write_output(self, markdown_text: str, output_path: str, format: str) -> bool:
        """Write markdown to target format."""
        format = format.lower()
        
        try:
            if format == 'html':
                return self._write_html(markdown_text, output_path)
            elif format == 'txt':
                return self._write_text(markdown_text, output_path)
            elif format == 'pdf':
                return self._write_pdf(markdown_text, output_path)
            elif format == 'docx':
                return self._write_docx(markdown_text, output_path)
            else:
                self.monitor.log_event('unsupported_output_format', {
                    'format': format
                }, severity='ERROR')
                return False
                
        except Exception as e:
            self.monitor.log_event('output_write_failed', {
                'format': format,
                'error': str(e)
            }, severity='ERROR')
            return False
    
    def _write_html(self, markdown_text: str, output_path: str) -> bool:
        """Convert Markdown to HTML."""
        try:
            import markdown
            
            # Convert markdown to HTML
            html_content = markdown.markdown(
                markdown_text,
                extensions=['extra', 'codehilite', 'tables', 'toc']
            )
            
            # Wrap in full HTML document
            html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>Converted Document</title>
<style>
body {{
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
    line-height: 1.6;
    color: #333;
    max-width: 900px;
    margin: 0 auto;
    padding: 20px;
}}
code {{
    background-color: #f6f8fa;
    padding: 2px 6px;
    border-radius: 3px;
    font-family: monospace;
}}
pre {{
    background-color: #f6f8fa;
    padding: 16px;
    border-radius: 6px;
    overflow: auto;
}}
table {{
    border-collapse: collapse;
    width: 100%;
    margin: 20px 0;
}}
th, td {{
    border: 1px solid #ddd;
    padding: 8px;
    text-align: left;
}}
th {{
    background-color: #f6f8fa;
}}
</style>
</head>
<body>
{html_content}
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html)
            
            return True
            
        except ImportError:
            self.monitor.log_event('dependency_missing', {
                'library': 'markdown',
                'install_command': 'pip install markdown'
            }, severity='ERROR')
            return False
        except Exception as e:
            self.monitor.log_event('html_conversion_failed', {
                'error': str(e)
            }, severity='ERROR')
            return False
    
    def _write_text(self, markdown_text: str, output_path: str) -> bool:
        """Write markdown as plain text (remove markdown syntax)."""
        # Simple markdown to text conversion
        # Remove markdown formatting
        import re
        
        text = markdown_text
        
        # Remove images: ![alt](url)
        text = re.sub(r'!\[([^\]]*)\]\([^\)]*\)', r'\1', text)
        
        # Remove links but keep text: [text](url)
        text = re.sub(r'\[([^\]]*)\]\([^\)]*\)', r'\1', text)
        
        # Remove bold/italic: **text** or *text*
        text = re.sub(r'\*\*([^\*]*)\*\*', r'\1', text)
        text = re.sub(r'\*([^\*]*)\*', r'\1', text)
        
        # Remove inline code: `code`
        text = re.sub(r'`([^`]*)`', r'\1', text)
        
        # Remove headings #
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        return True
    
    def _write_pdf(self, markdown_text: str, output_path: str) -> bool:
        """Convert Markdown to PDF (via HTML intermediate)."""
        try:
            # First convert to HTML
            import tempfile
            import markdown
            
            html_content = markdown.markdown(
                markdown_text,
                extensions=['extra', 'tables']
            )
            
            # Note: Actual PDF generation would require libraries like
            # pdfkit, weasyprint, or reportlab
            self.monitor.log_event('pdf_conversion_note', {
                'message': 'PDF conversion from Markdown requires additional setup',
                'suggestion': 'Use pandoc or specialized conversion tools'
            }, severity='WARNING')
            
            return False  # Not fully implemented
            
        except Exception as e:
            self.monitor.log_event('pdf_conversion_failed', {
                'error': str(e)
            }, severity='ERROR')
            return False
    
    def _write_docx(self, markdown_text: str, output_path: str) -> bool:
        """Convert Markdown to DOCX."""
        try:
            from docx import Document
            import re
            
            doc = Document()
            
            # Simple parsing - split by lines and interpret
            lines = markdown_text.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    continue
                
                # Headings
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)
            
            doc.save(output_path)
            return True
            
        except ImportError:
            self.monitor.log_event('dependency_missing', {
                'library': 'python-docx',
                'install_command': 'pip install python-docx'
            }, severity='ERROR')
            return False
        except Exception as e:
            self.monitor.log_event('docx_conversion_failed', {
                'error': str(e)
            }, severity='ERROR')
            return False
