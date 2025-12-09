"""
DOCX Converter

Handles Word document conversions.
Supported: DOCX/DOC ↔ TXT, MD, HTML, PDF
"""
import os
from typing import Optional, Dict
from src.converters.base_converter import BaseConverter


class DocxConverter(BaseConverter):
    """
    Handles Word document conversions.
    
    Supported: DOCX/DOC ↔ TXT, MD, HTML, PDF
    """
    
    def __init__(self):
        super().__init__()
        self.supported_input_formats = ['docx', 'doc']
        self.supported_output_formats = ['txt', 'md', 'html', 'pdf']
    
    def convert(self, input_path: str, output_path: str, output_format: str) -> bool:
        """Convert DOCX/DOC to target format."""
        timer_id = self.monitor.start_timer('docx_conversion')
        
        try:
            # Step 1: Validate
            self.monitor.log_event('conversion_start', {
                'input': input_path,
                'output_format': output_format,
                'output_path': output_path
            }, severity='INFO', step='Step 1/4: Validating input')
            
            if not self.validate_input(input_path):
                return False
            
            if not self.validate_output(output_path):
                return False
            
            # Step 2: Load document
            self.monitor.log_event('document_load', {
                'input': input_path
            }, severity='INFO', step='Step 2/4: Loading document')
            
            content = self._extract_content(input_path)
            
            if not content:
                return False
            
            # Step 3: Convert to target format
            self.monitor.log_event('format_conversion', {
                'output_format': output_format
            }, severity='INFO', step='Step 3/4: Converting to target format')
            
            result = self._write_output(content, output_path, output_format)
            
            # Step 4: Complete
            self.monitor.log_event('conversion_complete', {
                'input': input_path,
                'output': output_path,
                'success': result
            }, severity='INFO', step='Step 4/4: Conversion complete')
            
            return result
            
        except Exception as e:
            self.monitor.log_event('conversion_failed', {
                'input': input_path,
                'output_format': output_format,
                'error': str(e),
                'error_type': type(e).__name__
            }, severity='ERROR')
            return False
        finally:
            self.monitor.end_timer(timer_id)
    
    def _extract_content(self, docx_path: str) -> Optional[Dict]:
        """Extract content from DOCX/DOC file."""
        try:
            from docx import Document
            
            doc = Document(docx_path)
            
            content = {
                'paragraphs': [],
                'text': '',
                'tables': [],
                'properties': {}
            }
            
            # Extract core properties
            try:
                core_props = doc.core_properties
                content['properties'] = {
                    'title': core_props.title,
                    'author': core_props.author,
                    'subject': core_props.subject,
                    'created': str(core_props.created) if core_props.created else None,
                    'modified': str(core_props.modified) if core_props.modified else None
                }
            except:
                pass
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content['paragraphs'].append({
                        'text': text,
                        'style': para.style.name if para.style else None
                    })
                    content['text'] += text + '\n\n'
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                content['tables'].append(table_data)
            
            return content
            
        except ImportError:
            self.monitor.log_event('dependency_missing', {
                'library': 'python-docx',
                'install_command': 'pip install python-docx'
            }, severity='ERROR')
            return None
        except Exception as e:
            self.monitor.log_event('content_extraction_failed', {
                'file': docx_path,
                'error': str(e)
            }, severity='ERROR')
            return None
    
    def _write_output(self, content: Dict, output_path: str, format: str) -> bool:
        """Write content to target format."""
        format = format.lower()
        
        try:
            if format == 'txt':
                return self._write_text(content, output_path)
            elif format == 'md':
                return self._write_markdown(content, output_path)
            elif format == 'html':
                return self._write_html(content, output_path)
            elif format == 'pdf':
                return self._write_pdf(content, output_path)
            else:
                self.monitor.log_event('unsupported_output_format', {
                    'format': format
                }, severity='ERROR')
                return False
                
        except Exception as e:
            self.monitor.log_event('output_write_failed', {
                'format': format,
                'error': str(e)
            }, severity='ERROR')
            return False
    
    def _write_text(self, content: Dict, output_path: str) -> bool:
        """Write content as plain text."""
        with open(output_path, 'w', encoding='utf-8') as f:
            # Write properties
            props = content.get('properties', {})
            if props.get('title'):
                f.write(f"Title: {props['title']}\n\n")
            
            # Write text
            f.write(content['text'])
            
            # Write tables
            for i, table in enumerate(content.get('tables', [])):
                f.write(f"\n\n--- Table {i + 1} ---\n")
                for row in table:
                    f.write(' | '.join(row) + '\n')
        
        return True
    
    def _write_markdown(self, content: Dict, output_path: str) -> bool:
        """Write content as Markdown."""
        with open(output_path, 'w', encoding='utf-8') as f:
            # Write properties
            props = content.get('properties', {})
            if props.get('title'):
                f.write(f"# {props['title']}\n\n")
            
            if props.get('author'):
                f.write(f"**Author:** {props['author']}\n\n")
            
            # Write paragraphs with style hints
            for para in content.get('paragraphs', []):
                text = para['text']
                style = para.get('style', '')
                
                if 'Heading 1' in style:
                    f.write(f"# {text}\n\n")
                elif 'Heading 2' in style:
                    f.write(f"## {text}\n\n")
                elif 'Heading 3' in style:
                    f.write(f"### {text}\n\n")
                else:
                    f.write(f"{text}\n\n")
            
            # Write tables
            for table in content.get('tables', []):
                if table:
                    # Header row
                    f.write('| ' + ' | '.join(table[0]) + ' |\n')
                    f.write('|' + '|'.join(['---' for _ in table[0]]) + '|\n')
                    
                    # Data rows
                    for row in table[1:]:
                        f.write('| ' + ' | '.join(row) + ' |\n')
                    f.write('\n')
        
        return True
    
    def _write_html(self, content: Dict, output_path: str) -> bool:
        """Write content as HTML."""
        html = "<!DOCTYPE html>\n<html>\n<head>\n"
        html += "<meta charset=\"UTF-8\">\n"
        
        props = content.get('properties', {})
        title = props.get('title', 'Document')
        html += f"<title>{title}</title>\n"
        html += "<style>body { font-family: Arial, sans-serif; margin: 40px; }</style>\n"
        html += "</head>\n<body>\n"
        
        if props.get('title'):
            html += f"<h1>{props['title']}</h1>\n"
        
        # Metadata
        if props.get('author') or props.get('created'):
            html += "<div class='metadata'>\n"
            if props.get('author'):
                html += f"<p><strong>Author:</strong> {props['author']}</p>\n"
            if props.get('created'):
                html += f"<p><strong>Created:</strong> {props['created']}</p>\n"
            html += "</div>\n"
        
        # Content
        for para in content.get('paragraphs', []):
            style = para.get('style', '')
            text = para['text']
            
            if 'Heading 1' in style:
                html += f"<h1>{text}</h1>\n"
            elif 'Heading 2' in style:
                html += f"<h2>{text}</h2>\n"
            elif 'Heading 3' in style:
                html += f"<h3>{text}</h3>\n"
            else:
                html += f"<p>{text}</p>\n"
        
        # Tables
        for table in content.get('tables', []):
            html += "<table border='1' style='border-collapse: collapse; margin: 20px 0;'>\n"
            for i, row in enumerate(table):
                html += "<tr>\n"
                tag = "th" if i == 0 else "td"
                for cell in row:
                    html += f"<{tag} style='padding: 8px;'>{cell}</{tag}>\n"
                html += "</tr>\n"
            html += "</table>\n"
        
        html += "</body>\n</html>"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
        
        return True
    
    def _write_pdf(self, content: Dict, output_path: str) -> bool:
        """Write content as PDF (requires intermediate conversion)."""
        try:
            # Create a temporary DOCX first, then convert to PDF
            from docx import Document
            
            doc = Document()
            
            # Add title
            props = content.get('properties', {})
            if props.get('title'):
                doc.add_heading(props['title'], 0)
            
            # Add paragraphs
            for para in content.get('paragraphs', []):
                doc.add_paragraph(para['text'])
            
            # Add tables
            for table_data in content.get('tables', []):
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    for i, row in enumerate(table_data):
                        for j, cell_text in enumerate(row):
                            table.rows[i].cells[j].text = cell_text
            
            # Save temporary DOCX
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp_path = tmp.name
                doc.save(tmp_path)
            
            # Note: Actual PDF conversion would require additional libraries
            # like docx2pdf (Windows-only) or conversion services
            # For now, we'll note this limitation
            self.monitor.log_event('pdf_conversion_note', {
                'message': 'PDF conversion from DOCX requires additional setup',
                'suggestion': 'Use Microsoft Word, LibreOffice, or cloud conversion services'
            }, severity='WARNING')
            
            # Clean up temp file
            os.remove(tmp_path)
            
            return False  # Not fully implemented without external dependencies
            
        except Exception as e:
            self.monitor.log_event('pdf_conversion_failed', {
                'error': str(e)
            }, severity='ERROR')
            return False
